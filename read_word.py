from docx import Document
from baza import to_cyrillic, to_latin
from anaylise import has_cyrillic
def word_reader(file):
    document = Document(file)
    for para in document.paragraphs:
        inline = para.runs
        for i in range(len(inline)):
                if has_cyrillic(inline[i].text):
                    inline[i].text = to_latin(inline[i].text)
                else:
                    inline[i].text = to_cyrillic(text=inline[i].text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    inline = para.runs
                    for i in range(len(inline)):
                        if has_cyrillic(inline[i].text):
                            inline[i].text = to_latin(inline[i].text)
                        else:
                            inline[i].text = to_cyrillic(text=inline[i].text)
    document.save(file)
